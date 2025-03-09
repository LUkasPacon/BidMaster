#!/usr/bin/env python3
"""
Skript pro vytvoření vzorové šablony dokumentu.
"""
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Přidání nadřazeného adresáře do cesty pro import
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

def create_template():
    """
    Vytvoří vzorovou šablonu dokumentu.
    """
    # Vytvoření adresáře pro šablony, pokud neexistuje
    template_dir = Path("data/templates")
    template_dir.mkdir(parents=True, exist_ok=True)
    
    # Vytvoření adresáře pro generované dokumenty, pokud neexistuje
    generated_dir = Path("data/generated")
    generated_dir.mkdir(parents=True, exist_ok=True)
    
    # Cesta k šabloně
    template_path = template_dir / "template.docx"
    
    # Vytvoření dokumentu
    doc = Document()
    
    # Nastavení stylu nadpisu 1
    style_heading1 = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    style_heading1.base_style = doc.styles['Heading 1']
    font = style_heading1.font
    font.name = 'Arial'
    font.size = Pt(16)
    font.bold = True
    font.color.rgb = RGBColor(0, 51, 102)
    
    # Nastavení stylu nadpisu 2
    style_heading2 = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    style_heading2.base_style = doc.styles['Heading 2']
    font = style_heading2.font
    font.name = 'Arial'
    font.size = Pt(14)
    font.bold = True
    font.color.rgb = RGBColor(0, 51, 102)
    
    # Nastavení stylu normálního textu
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Nastavení stylu pro tabulky
    style_table = doc.styles.add_style('TableStyle', WD_STYLE_TYPE.PARAGRAPH)
    font = style_table.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Přidání titulní strany
    p = doc.add_paragraph("Nabídka implementace MidPoint")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph("pro {client_name}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    # Přidání mezery
    for _ in range(5):
        doc.add_paragraph()
    
    p = doc.add_paragraph("Datum: {date}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("Verze: {version}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Přidání obsahu
    doc.add_page_break()
    p = doc.add_paragraph("Obsah")
    p.style = 'CustomHeading1'
    doc.add_paragraph("Tento obsah bude automaticky vygenerován při otevření dokumentu.")
    
    # Přidání úvodu
    doc.add_page_break()
    p = doc.add_paragraph("1. Úvod")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph("{introduction}")
    
    # Přidání popisu řešení
    doc.add_page_break()
    p = doc.add_paragraph("2. Popis řešení")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph("{solution_description}")
    
    # Přidání rozsahu prací
    p = doc.add_paragraph("3. Rozsah prací")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph("{scope_of_work}")
    
    # Přidání harmonogramu
    p = doc.add_paragraph("4. Harmonogram")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph("{timeline}")
    
    # Přidání cenové nabídky
    p = doc.add_paragraph("5. Cenová nabídka")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph("{pricing}")
    
    # Přidání kontaktních informací
    doc.add_page_break()
    p = doc.add_paragraph("6. Kontaktní informace")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph("{contact_info}")
    
    # Uložení dokumentu
    doc.save(template_path)
    
    print(f"Šablona byla vytvořena a uložena na: {template_path}")
    print(f"Adresář pro generované dokumenty: {generated_dir}")

if __name__ == "__main__":
    create_template() 