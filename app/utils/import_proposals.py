#!/usr/bin/env python3
"""
Utilita pro import nabídek do vektorové databáze.
"""
import os
import sys
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional
import docx
import re
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from dotenv import load_dotenv
from langchain_pinecone import PineconeVectorStore
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    JSONLoader
)
import pinecone

# Přidání podpory pro PDF
try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        print("VAROVÁNÍ: Knihovna pro práci s PDF není nainstalována. Použijte 'pip install pypdf' nebo 'pip install PyPDF2'.")
        PdfReader = None

# Přidání kořenového adresáře do cesty pro import
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from app.utils.vector_store import (
    init_vector_store,
    add_documents_to_vector_store,
    get_embeddings
)
from app.config import get_config

config = get_config()

load_dotenv()

def init_pinecone() -> PineconeVectorStore:
    """Inicializace Pinecone."""
    pinecone.init(
        api_key=os.getenv("PINECONE_API_KEY", ""),
        environment=os.getenv("PINECONE_ENV", "")
    )
    return PineconeVectorStore(
        index_name=os.getenv("PINECONE_INDEX", "bidmaster"),
        namespace=os.getenv("PINECONE_NAMESPACE", "proposals")
    )

def load_document(file_path: str) -> List[Dict[str, Any]]:
    """
    Načte dokument podle typu souboru.
    
    Args:
        file_path: Cesta k souboru
        
    Returns:
        List[Dict[str, Any]]: Seznam dokumentů s metadaty
    """
    path = Path(file_path)
    
    if path.suffix.lower() == '.pdf':
        loader = PyPDFLoader(str(path))
        return loader.load()
    elif path.suffix.lower() == '.docx':
        loader = Docx2txtLoader(str(path))
        return loader.load()
    elif path.suffix.lower() == '.json':
        loader = JSONLoader(
            file_path=str(path),
            jq_schema='.content',
            text_content=False
        )
        return loader.load()
    else:
        raise ValueError(f"Nepodporovaný typ souboru: {path.suffix}")

def process_documents(documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Zpracuje dokumenty pro import.
    
    Args:
        documents: Seznam dokumentů
        
    Returns:
        List[Dict[str, Any]]: Zpracované dokumenty
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        length_function=len
    )
    
    processed_docs = []
    for doc in documents:
        chunks = text_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            processed_docs.append({
                "page_content": chunk,
                "metadata": {
                    **doc.metadata,
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                }
            })
    return processed_docs

def import_proposals(file_paths: List[str]) -> None:
    """
    Importuje návrhy do vektorové databáze.
    
    Args:
        file_paths: Seznam cest k souborům
    """
    vectorstore = init_pinecone()
    
    for file_path in file_paths:
        try:
            # Načtení dokumentu
            documents = load_document(file_path)
            
            # Zpracování dokumentů
            processed_docs = process_documents(documents)
            
            # Import do Pinecone
            vectorstore.add_documents(processed_docs)
            
            print(f"Úspěšně importován soubor: {file_path}")
            
        except Exception as e:
            print(f"Chyba při importu souboru {file_path}: {e}")
            raise

def extract_text_from_docx(file_path: str) -> str:
    """
    Extrahuje text z DOCX dokumentu.
    
    Args:
        file_path: Cesta k DOCX dokumentu
        
    Returns:
        str: Extrahovaný text
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extrakce textu z paragrafů
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extrakce textu z tabulek
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
        
        return "\n".join(full_text)
    except Exception as e:
        print(f"Chyba při extrakci textu z dokumentu {file_path}: {e}")
        return ""

def extract_metadata_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Extrahuje metadata z DOCX dokumentu.
    
    Args:
        file_path: Cesta k DOCX dokumentu
        
    Returns:
        Dict[str, Any]: Extrahovaná metadata
    """
    try:
        doc = docx.Document(file_path)
        metadata = {
            "source": file_path,
            "title": os.path.basename(file_path),
            "client_name": "",
            "date": "",
            "version": ""
        }
        
        # Pokus o extrakci názvu klienta
        client_pattern = re.compile(r"pro\s+(.+?)(?:\s+ze\s+dne|\s*$)")
        for para in doc.paragraphs[:20]:  # Prohledáme prvních 20 odstavců
            if "nabídka" in para.text.lower() and "pro" in para.text.lower():
                match = client_pattern.search(para.text)
                if match:
                    metadata["client_name"] = match.group(1).strip()
                    break
        
        # Pokus o extrakci data
        date_pattern = re.compile(r"\b(\d{1,2}\.\s*\d{1,2}\.\s*\d{4}|\d{4}-\d{2}-\d{2})\b")
        for para in doc.paragraphs[:30]:  # Prohledáme prvních 30 odstavců
            match = date_pattern.search(para.text)
            if match:
                metadata["date"] = match.group(1)
                break
        
        # Pokus o extrakci verze
        version_pattern = re.compile(r"verze\s*[:]\s*([0-9.]+)", re.IGNORECASE)
        for para in doc.paragraphs[:30]:  # Prohledáme prvních 30 odstavců
            match = version_pattern.search(para.text)
            if match:
                metadata["version"] = match.group(1)
                break
        
        return metadata
    except Exception as e:
        print(f"Chyba při extrakci metadat z dokumentu {file_path}: {e}")
        return {
            "source": file_path,
            "title": os.path.basename(file_path)
        }

def process_docx_file(file_path: str) -> Dict[str, Any]:
    """
    Zpracuje DOCX soubor a vrátí jeho obsah a metadata.
    
    Args:
        file_path: Cesta k DOCX souboru
        
    Returns:
        Slovník s textem a metadaty
    """
    try:
        # Maximální velikost textu v bajtech (50 KB)
        MAX_TEXT_SIZE = 50 * 1024
        
        text = extract_text_from_docx(file_path)
        metadata = extract_metadata_from_docx(file_path)
        
        # Přidání cesty k souboru do metadat
        metadata["source"] = file_path
        
        # Kontrola velikosti textu
        if len(text.encode('utf-8')) > MAX_TEXT_SIZE:
            print(f"Varování: Text z {file_path} je příliš velký. Bude zkrácen.")
            text = text[:MAX_TEXT_SIZE // 4]  # Přibližně 50 KB v UTF-8
            metadata["truncated_text"] = True
        
        return {
            "text": text,
            "metadata": metadata
        }
    except Exception as e:
        print(f"Chyba při zpracování DOCX souboru {file_path}: {e}")
        return {
            "text": f"Chyba při zpracování souboru: {e}",
            "metadata": {"source": file_path, "error": str(e)}
        }

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extrahuje text z PDF dokumentu.
    
    Args:
        file_path: Cesta k PDF dokumentu
        
    Returns:
        str: Extrahovaný text
    """
    if PdfReader is None:
        print(f"Nelze extrahovat text z PDF dokumentu {file_path}: Knihovna pro práci s PDF není nainstalována.")
        return ""
    
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        # Extrakce textu ze všech stránek
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"Chyba při extrakci textu z PDF dokumentu {file_path}: {e}")
        return ""

def extract_metadata_from_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extrahuje metadata z PDF dokumentu.
    
    Args:
        file_path: Cesta k PDF dokumentu
        
    Returns:
        Dict[str, Any]: Extrahovaná metadata
    """
    if PdfReader is None:
        print(f"Nelze extrahovat metadata z PDF dokumentu {file_path}: Knihovna pro práci s PDF není nainstalována.")
        return {
            "source": file_path,
            "title": os.path.basename(file_path)
        }
    
    try:
        reader = PdfReader(file_path)
        metadata = {
            "source": file_path,
            "title": os.path.basename(file_path),
            "client_name": "",
            "date": "",
            "version": ""
        }
        
        # Extrakce metadat z PDF
        pdf_info = reader.metadata
        if pdf_info:
            if pdf_info.title:
                metadata["title"] = pdf_info.title
            if pdf_info.author:
                metadata["author"] = pdf_info.author
            if pdf_info.subject:
                metadata["subject"] = pdf_info.subject
            if pdf_info.creator:
                metadata["creator"] = pdf_info.creator
            if pdf_info.producer:
                metadata["producer"] = pdf_info.producer
            if pdf_info.creation_date:
                metadata["creation_date"] = str(pdf_info.creation_date)
        
        # Extrakce textu pro další analýzu
        text = extract_text_from_pdf(file_path)
        
        # Pokus o extrakci názvu klienta
        client_pattern = re.compile(r"pro\s+(.+?)(?:\s+ze\s+dne|\s*$)")
        for line in text.split('\n')[:30]:  # Prohledáme prvních 30 řádků
            if "nabídka" in line.lower() and "pro" in line.lower():
                match = client_pattern.search(line)
                if match:
                    metadata["client_name"] = match.group(1).strip()
                    break
        
        # Pokus o extrakci data
        date_pattern = re.compile(r"\b(\d{1,2}\.\s*\d{1,2}\.\s*\d{4}|\d{4}-\d{2}-\d{2})\b")
        for line in text.split('\n')[:50]:  # Prohledáme prvních 50 řádků
            match = date_pattern.search(line)
            if match:
                metadata["date"] = match.group(1)
                break
        
        # Pokus o extrakci verze
        version_pattern = re.compile(r"verze\s*[:]\s*([0-9.]+)", re.IGNORECASE)
        for line in text.split('\n')[:50]:  # Prohledáme prvních 50 řádků
            match = version_pattern.search(line)
            if match:
                metadata["version"] = match.group(1)
                break
        
        return metadata
    except Exception as e:
        print(f"Chyba při extrakci metadat z PDF dokumentu {file_path}: {e}")
        return {
            "source": file_path,
            "title": os.path.basename(file_path)
        }

def process_pdf_file(file_path: str) -> Dict[str, Any]:
    """
    Zpracuje PDF soubor a vrátí jeho obsah a metadata.
    
    Args:
        file_path: Cesta k PDF souboru
        
    Returns:
        Slovník s textem a metadaty
    """
    try:
        # Maximální velikost textu v bajtech (50 KB)
        MAX_TEXT_SIZE = 50 * 1024
        
        text = extract_text_from_pdf(file_path)
        metadata = extract_metadata_from_pdf(file_path)
        
        # Přidání cesty k souboru do metadat
        metadata["source"] = file_path
        
        # Kontrola velikosti textu
        if len(text.encode('utf-8')) > MAX_TEXT_SIZE:
            print(f"Varování: Text z {file_path} je příliš velký. Bude zkrácen.")
            text = text[:MAX_TEXT_SIZE // 4]  # Přibližně 50 KB v UTF-8
            metadata["truncated_text"] = True
        
        return {
            "text": text,
            "metadata": metadata
        }
    except Exception as e:
        print(f"Chyba při zpracování PDF souboru {file_path}: {e}")
        return {
            "text": f"Chyba při zpracování souboru: {e}",
            "metadata": {"source": file_path, "error": str(e)}
        }

def process_json_file(file_path: str) -> Dict[str, Any]:
    """
    Zpracuje JSON soubor a vrátí jeho obsah a metadata.
    
    Args:
        file_path: Cesta k JSON souboru
        
    Returns:
        Slovník s textem a metadaty
    """
    try:
        # Maximální velikost textu v bajtech (50 KB)
        MAX_TEXT_SIZE = 50 * 1024
        
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        # Kontrola, zda JSON obsahuje požadované klíče
        if "text" not in data:
            raise ValueError("JSON neobsahuje klíč 'text'")
        
        text = data["text"]
        
        # Kontrola velikosti textu
        if len(text.encode('utf-8')) > MAX_TEXT_SIZE:
            print(f"Varování: Text z {file_path} je příliš velký. Bude zkrácen.")
            text = text[:MAX_TEXT_SIZE // 4]  # Přibližně 50 KB v UTF-8
            data["metadata"] = data.get("metadata", {})
            data["metadata"]["truncated_text"] = True
        
        # Pokud JSON neobsahuje metadata, vytvoříme prázdný slovník
        if "metadata" not in data:
            data["metadata"] = {}
        
        # Přidání cesty k souboru do metadat
        data["metadata"]["source"] = file_path
        
        return {
            "text": text,
            "metadata": data["metadata"]
        }
    except Exception as e:
        print(f"Chyba při zpracování JSON souboru {file_path}: {e}")
        return {
            "text": f"Chyba při zpracování souboru: {e}",
            "metadata": {"source": file_path, "error": str(e)}
        }

def process_directory(directory_path: str) -> List[Dict[str, Any]]:
    """
    Zpracuje adresář s nabídkami.
    
    Args:
        directory_path: Cesta k adresáři
        
    Returns:
        List[Dict[str, Any]]: Seznam zpracovaných dokumentů
    """
    documents = []
    
    # Kontrola, zda adresář existuje
    if not os.path.exists(directory_path):
        print(f"Adresář {directory_path} neexistuje")
        return documents
    
    # Procházení souborů v adresáři
    for root, _, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)
            
            # Zpracování podle typu souboru
            if file.lower().endswith('.docx'):
                doc = process_docx_file(file_path)
                if doc["text"]:
                    documents.append(doc)
            elif file.lower().endswith('.json'):
                doc = process_json_file(file_path)
                if doc["text"]:
                    documents.append(doc)
            elif file.lower().endswith('.pdf'):
                doc = process_pdf_file(file_path)
                if doc["text"]:
                    documents.append(doc)
    
    return documents

def main():
    """
    Hlavní funkce pro spuštění importu nabídek.
    """
    parser = argparse.ArgumentParser(description="Import nabídek do vektorové databáze")
    parser.add_argument("directory", help="Cesta k adresáři s nabídkami")
    parser.add_argument("--namespace", help="Namespace pro vektorovou databázi", default=None)
    
    args = parser.parse_args()
    
    import_proposals(args.directory, args.namespace)

if __name__ == "__main__":
    main() 