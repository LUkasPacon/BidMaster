"""
Utilita pro generování DOCX dokumentů.
"""
import os
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_proposal_document(
    data: Dict[str, Any],
    template_path: Optional[str] = None,
    output_path: Optional[str] = None
) -> str:
    """
    Vytvoří dokument nabídky na základě šablony a dat.
    
    Args:
        data: Data pro vyplnění šablony
        template_path: Cesta k šabloně (volitelné)
        output_path: Cesta pro uložení výsledného dokumentu (volitelné)
        
    Returns:
        str: Cesta k vytvořenému dokumentu
    """
    # Pokud je zadána cesta k šabloně a šablona existuje, použijeme ji
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        # Jinak vytvoříme nový dokument
        doc = Document()
        setup_document_styles(doc)
    
    # Vyplnění dokumentu daty
    fill_document_with_data(doc, data)
    
    # Pokud není zadána cesta pro uložení, vytvoříme ji
    if not output_path:
        client_name = data.get("client_name", "klient").replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"data/generated/nabidka_{client_name}_{timestamp}.docx"
        
        # Ujistíme se, že adresář existuje
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Uložení dokumentu
    try:
        doc.save(output_path)
        print(f"Dokument byl úspěšně uložen na: {output_path}")
    except Exception as e:
        print(f"Chyba při ukládání dokumentu: {e}")
        # Zkusíme uložit do aktuálního adresáře
        fallback_path = f"nabidka_{data.get('client_name', 'klient').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        try:
            doc.save(fallback_path)
            print(f"Dokument byl uložen do aktuálního adresáře: {fallback_path}")
            output_path = fallback_path
        except Exception as e2:
            print(f"Nelze uložit dokument ani do aktuálního adresáře: {e2}")
            return "Chyba při ukládání dokumentu"
    
    return output_path

def setup_document_styles(doc: Document) -> None:
    """
    Nastaví styly dokumentu.
    
    Args:
        doc: Dokument
    """
    # Nastavení stylu nadpisu 1
    style_heading1 = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    style_heading1.base_style = doc.styles['Heading 1']
    font = style_heading1.font
    font.name = 'Arial'
    font.size = Pt(16)
    font.bold = True
    font.color.rgb = RGBColor(0, 51, 102)
    
    # Nastavení stylu nadpisu 2
    style_heading2 = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    style_heading2.base_style = doc.styles['Heading 2']
    font = style_heading2.font
    font.name = 'Arial'
    font.size = Pt(14)
    font.bold = True
    font.color.rgb = RGBColor(0, 51, 102)
    
    # Nastavení stylu normálního textu
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Nastavení stylu pro tabulky
    style_table = doc.styles.add_style('TableStyle', WD_STYLE_TYPE.PARAGRAPH)
    font = style_table.font
    font.name = 'Arial'
    font.size = Pt(10)

def fill_document_with_data(doc: Document, data: Dict[str, Any]) -> None:
    """
    Vyplní dokument daty.
    
    Args:
        doc: Dokument
        data: Data pro vyplnění
    """
    # Přidání titulní strany
    add_title_page(doc, data)
    
    # Přidání obsahu
    doc.add_page_break()
    p = doc.add_paragraph("Obsah")
    p.style = 'CustomHeading1'
    doc.add_paragraph("Tento obsah bude automaticky vygenerován při otevření dokumentu.")
    
    # Přidání úvodu
    doc.add_page_break()
    p = doc.add_paragraph("1. Úvod")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph(data.get("introduction", ""))
    
    # Přidání popisu řešení
    doc.add_page_break()
    p = doc.add_paragraph("2. Popis řešení")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph(data.get("solution_description", ""))
    
    # Přidání rozsahu prací
    p = doc.add_paragraph("3. Rozsah prací")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph(data.get("scope_of_work", ""))
    
    # Přidání harmonogramu
    p = doc.add_paragraph("4. Harmonogram")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph(data.get("timeline", ""))
    
    # Přidání cenové nabídky
    p = doc.add_paragraph("5. Cenová nabídka")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph(data.get("pricing", ""))
    
    # Přidání kontaktních informací
    doc.add_page_break()
    p = doc.add_paragraph("6. Kontaktní informace")
    p.style = 'CustomHeading1'
    
    p = doc.add_paragraph(data.get("contact_info", ""))

def add_title_page(doc: Document, data: Dict[str, Any]) -> None:
    """
    Přidá titulní stranu do dokumentu.
    
    Args:
        doc: Dokument
        data: Data pro vyplnění
    """
    # Přidání loga (pokud existuje)
    logo_path = data.get("logo_path")
    if logo_path and os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(2.0))
    
    # Přidání mezery
    for _ in range(5):
        doc.add_paragraph()
    
    # Přidání názvu nabídky
    p = doc.add_paragraph("Nabídka implementace MidPoint")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Přidání názvu klienta
    p = doc.add_paragraph(f"pro {data.get('client_name', 'klienta')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    # Přidání mezery
    for _ in range(5):
        doc.add_paragraph()
    
    # Přidání data
    p = doc.add_paragraph(f"Datum: {data.get('date', datetime.now().strftime('%d.%m.%Y'))}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Přidání verze
    p = doc.add_paragraph(f"Verze: {data.get('version', '1.0')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER 