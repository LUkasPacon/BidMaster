import streamlit as st
from pathlib import Path
import json
from typing import Dict, List, Optional
import sys
import os
from datetime import datetime
from docx import Document
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage, AIMessage

# Přidání root adresáře do PYTHONPATH
root_dir = str(Path(__file__).parent.parent.parent)
if root_dir not in sys.path:
    sys.path.append(root_dir)

from app.utils.import_proposals import import_proposals, load_document
from app.utils.search_proposals import search_proposals
from app.proposal_graph import SimpleStateGraph, Step, ProposalState

# Konfigurace stránky
st.set_page_config(
    page_title="BidMaster",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Styly
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
        margin-bottom: 1rem;
    }
    .status-box {
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
    }
    .user-message {
        background-color: #2e3136;
    }
    .assistant-message {
        background-color: #3a3f45;
    }
    .system-message {
        background-color: #1e2124;
        color: #8e9297;
        font-style: italic;
    }
    </style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Inicializace session state pro uchování stavu aplikace."""
    if 'current_step' not in st.session_state:
        st.session_state.current_step = Step.ANALYZE_REQUEST
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'proposal_state' not in st.session_state:
        st.session_state.proposal_state = ProposalState()
    if 'graph' not in st.session_state:
        st.session_state.graph = SimpleStateGraph()
    if 'llm' not in st.session_state:
        st.session_state.llm = ChatOpenAI(temperature=0.7)
    if 'log' not in st.session_state:
        st.session_state.log = []

def add_log(message: str):
    """Přidá zprávu do logu."""
    timestamp = datetime.now().strftime("%H:%M:%S")
    st.session_state.log.append(f"[{timestamp}] {message}")

def render_chat_message(message: dict):
    """Vykreslí zprávu v chatu."""
    role = message["role"]
    content = message["content"]
    
    if role == "user":
        st.markdown(f'<div class="chat-message user-message">👤 **Vy**: {content}</div>', unsafe_allow_html=True)
    elif role == "assistant":
        st.markdown(f'<div class="chat-message assistant-message">🤖 **Asistent**: {content}</div>', unsafe_allow_html=True)
    elif role == "system":
        st.markdown(f'<div class="chat-message system-message">ℹ️ {content}</div>', unsafe_allow_html=True)

def get_next_question(state: ProposalState) -> str:
    """Získá další otázku od chatbota."""
    messages = [
        SystemMessage(content="""Jsi asistent pro analýzu poptávek na implementaci MidPoint.
        Tvým úkolem je získat všechny potřebné informace pro vytvoření nabídky.
        Ptej se stručně a jasně. Zaměř se na:
        - Technické požadavky
        - Integrační požadavky
        - Časové požadavky
        - Rozpočtové omezení
        - Specifické požadavky klienta"""),
    ]
    
    # Přidání historie konverzace
    for msg in state.chat_history:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        elif msg["role"] == "assistant":
            messages.append(AIMessage(content=msg["content"]))
    
    # Přidání kontextu aktuálního stavu
    context = f"""
    Dosud známé informace:
    - Počet uživatelů: {state.expected_users}
    - Termín dokončení: {state.deadline}
    - Požadované moduly: {', '.join(state.modules)}
    - Integrace: {', '.join(state.integrations) if state.integrations else 'nezadáno'}
    - Školení: {', '.join(state.training) if state.training else 'nezadáno'}
    - Úroveň podpory: {state.support_level}
    """
    messages.append(HumanMessage(content=f"Na základě těchto informací: {context}\nJaká je tvá další otázka?"))
    
    response = st.session_state.llm.invoke(messages)
    return response.content

def create_proposal_document(state: ProposalState) -> str:
    """Vytvoří dokument s nabídkou na základě získaných informací a uloží jej do adresáře generated_proposals."""
    
    try:
        # Inicializace dokumentu
        doc = Document()
        
        # Nadpis
        doc.add_heading('Nabídka implementace MidPoint', level=0)
        
        # Datum
        now = datetime.now()
        date_str = now.strftime("%d.%m.%Y")
        doc.add_paragraph(f'Datum: {date_str}')
        
        # Hledání relevantního kontextu
        context = ""
        if state.client_request:
            similar_proposals = search_proposals(state.client_request)
            if similar_proposals:
                context = "\n".join([p["content"] for p in similar_proposals[:3]])
        
        # Základní systémový prompt pro všechny sekce
        base_system_prompt = """Jsi profesionální business konzultant s rozsáhlými znalostmi v oblasti implementace MidPoint.
        Tvým úkolem je vytvořit profesionální část obchodní nabídky, která bude:
        - Věcná a konkrétní
        - Profesionální a formální
        - V českém jazyce
        - Bez technického žargonu, srozumitelná pro business stakeholdery
        - Přehledně strukturovaná
        
        DŮLEŽITÉ: Tvůj výstup MUSÍ být čistý text s běžným formátováním (odstavce, odrážky), 
        NE kód, NE JSON, NE dictionary ani jiné datové struktury!
        """
        
        # ===== SEKCE 1: POPIS ŘEŠENÍ =====
        doc.add_heading('2. Popis řešení', level=1)
        solution_prompt = f"""Vytvoř profesionální "Popis řešení" pro nabídku implementace MidPoint.
        
        Informace:
        - Poptávka klienta: {state.client_request}
        - Počet uživatelů: {state.expected_users}
        - Požadované moduly: {', '.join(state.modules)}
        - Integrace: {', '.join(state.integrations) if state.integrations else 'Není specifikováno'}
        
        Kontext z podobných nabídek:
        {context}
        
        Vytvoř stručný, ale informativní popis MidPoint řešení, který:
        1. Vysvětlí co je MidPoint a jaké jsou jeho hlavní přínosy
        2. Zdůrazní klíčové funkce relevantní pro tohoto klienta
        3. Popíše, jak řešení adresuje klientovy potřeby z poptávky
        
        DŮLEŽITÉ: Výstup formátuj jako běžný text s odstavci, NE jako JSON nebo Python dictionary!
        """
        
        messages = [
            SystemMessage(content=base_system_prompt),
            HumanMessage(content=solution_prompt)
        ]
        solution_response = st.session_state.llm.invoke(messages)
        doc.add_paragraph(solution_response.content)
        
        # ===== SEKCE 2: ROZSAH PRACÍ =====
        doc.add_heading('3. Rozsah prací', level=1)
        scope_prompt = f"""Vytvoř profesionální sekci "Rozsah prací" pro nabídku implementace MidPoint.
        
        Informace:
        - Poptávka klienta: {state.client_request}
        - Požadované moduly: {', '.join(state.modules)}
        - Integrace: {', '.join(state.integrations) if state.integrations else 'Není specifikováno'}
        - Školení: {', '.join(state.training) if state.training else 'Není specifikováno'}
        
        Kontext z podobných nabídek:
        {context}
        
        Vytvoř přehledně strukturovaný popis rozsahu prací, který:
        1. Začíná úvodním odstavcem shrnujícím celkový rozsah implementace
        2. Obsahuje seznam konkrétních oblastí práce (5-6 položek)
        3. Každou oblast detailně popíše včetně očekávaných výstupů
        
        Formátuj výstup jako souvislý text s nadpisy a odstavci. 
        DŮLEŽITÉ: Výstup MUSÍ být čistý text, NE kód, NE JSON struktura, NE Python dictionary!
        """
        
        messages = [
            SystemMessage(content=base_system_prompt),
            HumanMessage(content=scope_prompt)
        ]
        scope_response = st.session_state.llm.invoke(messages)
        doc.add_paragraph(scope_response.content)
        
        # ===== SEKCE 3: HARMONOGRAM =====
        doc.add_heading('4. Harmonogram', level=1)
        timeline_prompt = f"""Vytvoř profesionální sekci "Harmonogram" pro nabídku implementace MidPoint.
        
        Informace:
        - Poptávka klienta: {state.client_request}
        - Požadovaný termín dokončení: {state.deadline}
        - Rozsah prací: {', '.join(state.modules)}
        - Integrace: {', '.join(state.integrations) if state.integrations else 'Není specifikováno'}
        
        Kontext z podobných nabídek:
        {context}
        
        Vytvoř přehledný harmonogram implementace, který:
        1. Začíná úvodním odstavcem shrnujícím celkovou dobu implementace
        2. Obsahuje seznam jednotlivých fází (analýza, návrh, implementace, testování, atd.)
        3. Pro každou fázi uvádí její název, dobu trvání a stručný popis aktivit
        4. Končí závěrečným odstavcem shrnujícím celkový časový plán
        
        Formátuj výstup jako souvislý text s nadpisy, odstavci a případně s přehlednou tabulkou fází.
        DŮLEŽITÉ: Výstup MUSÍ být čistý text, NE kód, NE JSON struktura, NE Python dictionary!
        """
        
        messages = [
            SystemMessage(content=base_system_prompt),
            HumanMessage(content=timeline_prompt)
        ]
        timeline_response = st.session_state.llm.invoke(messages)
        doc.add_paragraph(timeline_response.content)
        
        # ===== SEKCE 4: CENOVÁ NABÍDKA =====
        doc.add_heading('5. Cenová nabídka', level=1)
        pricing_prompt = f"""Vytvoř profesionální sekci "Cenová nabídka" pro implementaci MidPoint.
        
        Informace:
        - Poptávka klienta: {state.client_request}
        - Počet uživatelů: {state.expected_users}
        - Požadované moduly: {', '.join(state.modules)}
        - Integrace: {', '.join(state.integrations) if state.integrations else 'Není specifikováno'}
        - Školení: {', '.join(state.training) if state.training else 'Není specifikováno'}
        
        Kontext z podobných nabídek:
        {context}
        
        Vytvoř detailní cenovou nabídku, která:
        1. Začíná úvodním odstavcem shrnujícím cenovou nabídku
        2. Obsahuje seznam položek s cenami (v Kč bez DPH)
        3. Pro každou položku uvádí její název a cenu
        4. Končí součtem s celkovou cenou
        
        Formátuj výstup jako souvislý text s nadpisy, odstavci a přehlednou tabulkou cen.
        DŮLEŽITÉ: Výstup MUSÍ být čistý text, NE kód, NE JSON struktura, NE Python dictionary!
        """
        
        messages = [
            SystemMessage(content=base_system_prompt),
            HumanMessage(content=pricing_prompt)
        ]
        pricing_response = st.session_state.llm.invoke(messages)
        doc.add_paragraph(pricing_response.content)
        
        # Uložení dokumentu
        # Vytvoření složky pro vygenerované nabídky, pokud neexistuje
        proposals_dir = Path("generated_proposals")
        if not proposals_dir.exists():
            proposals_dir.mkdir(parents=True)
        
        # Uložení dokumentu s časovým razítkem
        timestamp = now.strftime("%Y%m%d_%H%M%S")
        doc_path = proposals_dir / f"nabidka_midpoint_{timestamp}.docx"
        doc.save(str(doc_path))
        
        add_log(f"Vygenerována nabídka: {doc_path}")
        return str(doc_path)
    
    except Exception as e:
        st.error(f"Chyba při generování dokumentu: {str(e)}")
        add_log(f"Chyba při generování dokumentu: {str(e)}")
        return None

def render_sidebar():
    """Vykreslení postranního panelu."""
    with st.sidebar:
        st.title("BidMaster")
        st.markdown("---")
        
        # Navigační menu
        st.subheader("Menu")
        if st.button("Nová nabídka"):
            st.session_state.current_step = Step.ANALYZE_REQUEST
            st.session_state.chat_history = []
            st.session_state.proposal_state = ProposalState()
            add_log("Zahájena nová nabídka")
        
        if st.button("Import nabídek"):
            st.session_state.current_step = "IMPORT"
            add_log("Přepnuto na import nabídek")
        
        if st.button("Vyhledávání"):
            st.session_state.current_step = "SEARCH"
            add_log("Přepnuto na vyhledávání")
        
        st.markdown("---")
        
        # Status
        st.subheader("Status")
        st.write(f"Aktuální krok: {st.session_state.current_step}")
        
        # Log
        st.subheader("Log")
        for log_message in st.session_state.log:
            st.text(log_message)

def render_analyze_request():
    """Vykreslení formuláře pro analýzu poptávky."""
    st.header("Analýza poptávky")
    
    # Přidání možnosti nahrát soubor s poptávkou
    uploaded_file = st.file_uploader(
        "Nahrajte soubor s poptávkou (volitelné)",
        type=['pdf', 'docx', 'txt', 'json'],
        help="Můžete nahrát poptávku jako PDF, DOCX, TXT nebo JSON soubor."
    )
    
    # Pokud byl nahrán soubor, načteme jeho obsah
    if uploaded_file:
        try:
            # Vytvoření temp adresáře, pokud neexistuje
            temp_dir = Path("temp")
            temp_dir.mkdir(exist_ok=True)
            
            # Uložení souboru
            file_path = temp_dir / uploaded_file.name
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            add_log(f"Nahrán soubor: {uploaded_file.name}")
            
            # Načtení obsahu podle typu souboru
            if file_path.suffix.lower() == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    file_content = f.read()
            elif file_path.suffix.lower() in ['.pdf', '.docx']:
                documents = load_document(str(file_path))
                file_content = "\n\n".join([doc.page_content for doc in documents])
            elif file_path.suffix.lower() == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    file_content = data.get('text', '') if isinstance(data, dict) else str(data)
            
            # Smazání dočasného souboru
            file_path.unlink()
            
            st.success("Soubor byl úspěšně načten!")
            add_log("Soubor byl úspěšně zpracován")
            
        except Exception as e:
            st.error(f"Chyba při načítání souboru: {str(e)}")
            add_log(f"Chyba při zpracování souboru: {str(e)}")
            file_content = ""
    else:
        file_content = ""
    
    with st.form("analyze_request_form"):
        client_request = st.text_area(
            "Zadejte poptávku klienta",
            value=file_content,
            height=200,
            placeholder="Popište požadavky klienta na implementaci MidPoint..."
        )
        
        col1, col2 = st.columns(2)
        with col1:
            expected_users = st.number_input("Počet uživatelů", min_value=1, value=100)
        with col2:
            deadline = st.date_input("Požadovaný termín dokončení")
        
        modules = st.multiselect(
            "Požadované moduly",
            ["Role Management", "Workflow Management", "Self-service", "Reporting", "Audit", "Rekonciliace"],
            default=["Role Management"]
        )
        
        submit = st.form_submit_button("Analyzovat poptávku")
        
        if submit and client_request:
            st.session_state.proposal_state.client_request = client_request
            st.session_state.proposal_state.expected_users = expected_users
            st.session_state.proposal_state.deadline = deadline
            st.session_state.proposal_state.modules = modules
            st.session_state.current_step = Step.GATHER_INFORMATION
            
            # Přidání první zprávy do historie
            st.session_state.chat_history.append({
                "role": "user",
                "content": client_request
            })
            
            add_log("Poptávka byla analyzována")

def render_gather_information():
    """Vykreslení formuláře pro sběr doplňujících informací."""
    st.header("Doplňující informace")
    
    # Zobrazení historie konverzace
    for message in st.session_state.chat_history:
        render_chat_message(message)
    
    # Získání další otázky od chatbota
    if not st.session_state.chat_history or st.session_state.chat_history[-1]["role"] == "user":
        next_question = get_next_question(st.session_state.proposal_state)
        st.session_state.chat_history.append({
            "role": "assistant",
            "content": next_question
        })
        render_chat_message(st.session_state.chat_history[-1])
    
    # Formulář pro odpověď
    with st.form("gather_info_form"):
        user_response = st.text_area("Vaše odpověď", height=100)
        col1, col2 = st.columns([1, 4])
        
        with col1:
            submit = st.form_submit_button("Odpovědět")
        with col2:
            if st.form_submit_button("Přejít na generování nabídky"):
                st.session_state.current_step = Step.GENERATE_PROPOSAL
                add_log("Přechod na generování nabídky")
                return
        
        if submit and user_response:
            st.session_state.chat_history.append({
                "role": "user",
                "content": user_response
            })
            add_log("Přidána odpověď uživatele")
            st.rerun()

def render_generate_proposal():
    """Vykreslení generování nabídky."""
    st.header("Generování nabídky")
    
    # Zobrazení shrnutí
    st.subheader("Shrnutí požadavků")
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("**Základní informace**")
        st.write(f"Počet uživatelů: {st.session_state.proposal_state.expected_users}")
        st.write(f"Termín dokončení: {st.session_state.proposal_state.deadline}")
        st.write("\n**Moduly**")
        for module in st.session_state.proposal_state.modules:
            st.write(f"- {module}")
    
    with col2:
        st.write("**Integrace**")
        for integration in st.session_state.proposal_state.integrations:
            st.write(f"- {integration}")
        st.write("\n**Školení**")
        for training in st.session_state.proposal_state.training:
            st.write(f"- {training}")
        st.write(f"\n**Podpora**: {st.session_state.proposal_state.support_level}")
    
    if st.button("Generovat nabídku"):
        try:
            file_path = create_proposal_document(st.session_state.proposal_state)
            st.success(f"Nabídka byla vygenerována a uložena do: {file_path}")
            add_log(f"Nabídka byla vygenerována: {file_path}")
            
            # Nabídnutí stažení souboru
            with open(file_path, "rb") as file:
                st.download_button(
                    label="Stáhnout nabídku",
                    data=file,
                    file_name=os.path.basename(file_path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"Chyba při generování nabídky: {str(e)}")
            add_log(f"Chyba při generování nabídky: {str(e)}")

def render_import():
    """Vykreslení importu nabídek."""
    st.header("Import nabídek")
    
    uploaded_files = st.file_uploader(
        "Vyberte soubory k importu",
        accept_multiple_files=True,
        type=['pdf', 'docx', 'json']
    )
    
    if uploaded_files:
        # Vytvoření temp adresáře, pokud neexistuje
        temp_dir = Path("temp")
        temp_dir.mkdir(exist_ok=True)
        
        for file in uploaded_files:
            file_path = temp_dir / file.name
            with open(file_path, "wb") as f:
                f.write(file.getbuffer())
            
            try:
                import_proposals([str(file_path)])
                st.success(f"Soubor {file.name} byl úspěšně importován")
                add_log(f"Importován soubor: {file.name}")
            except Exception as e:
                st.error(f"Chyba při importu souboru {file.name}: {str(e)}")
                add_log(f"Chyba při importu souboru {file.name}: {str(e)}")
            finally:
                if file_path.exists():
                    file_path.unlink()

def render_search():
    """Vykreslení vyhledávání v nabídkách."""
    st.header("Vyhledávání v nabídkách")
    
    query = st.text_input("Zadejte hledaný text")
    
    if query:
        try:
            add_log(f"Vyhledávání: {query}")
            results = search_proposals(query)
            
            if results:
                for i, result in enumerate(results, 1):
                    with st.expander(f"Výsledek {i}"):
                        st.write(result.page_content)
                        st.write("---")
                        st.write("**Metadata:**")
                        st.json(result.metadata)
                add_log(f"Nalezeno {len(results)} výsledků")
            else:
                st.info("Nebyly nalezeny žádné výsledky.")
                add_log("Žádné výsledky")
                
        except Exception as e:
            st.error(f"Chyba při vyhledávání: {str(e)}")
            add_log(f"Chyba při vyhledávání: {str(e)}")

def main():
    """Hlavní funkce aplikace."""
    initialize_session_state()
    render_sidebar()
    
    # Vykreslení hlavního obsahu podle aktuálního kroku
    if st.session_state.current_step == Step.ANALYZE_REQUEST:
        render_analyze_request()
    elif st.session_state.current_step == Step.GATHER_INFORMATION:
        render_gather_information()
    elif st.session_state.current_step == Step.GENERATE_PROPOSAL:
        render_generate_proposal()
    elif st.session_state.current_step == "IMPORT":
        render_import()
    elif st.session_state.current_step == "SEARCH":
        render_search()

if __name__ == "__main__":
    main() 